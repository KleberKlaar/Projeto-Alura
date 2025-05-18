# Importações necessárias
import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog
import threading
import os
import warnings
# Importações do seu código original para os agentes
from google.adk.agents import Agent
from google.adk.runners import Runner
from google.adk.sessions import InMemorySessionService
from google.adk.tools import google_search
from google.genai import types
from google import genai

# Importação para criar documentos Word
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

warnings.filterwarnings("ignore")

os.environ["GOOGLE_API_KEY"] = 'INSIRA SUA API_KEY' # <<< SUBSTITUA PELA SUA CHAVE REAL

# Configura o cliente da SDK do Gemini
client = genai.Client()
MODEL_ID = "gemini-2.0-flash" # Mantendo o modelo original

# --- Função auxiliar que envia uma mensagem para um agente via Runner e retorna a resposta final ---
def call_agent(agent: Agent, message_text: str) -> str:
    """
    Envia uma mensagem para um agente via Runner e retorna a resposta final.
    Inclui tratamento para respostas inesperadas da API.
    """
    session_service = InMemorySessionService()
    session = session_service.create_session(app_name=agent.name, user_id="user1", session_id="session1")
    runner = Runner(agent=agent, app_name=agent.name, session_service=session_service)
    content = types.Content(role="user", parts=[types.Part(text=message_text)])

    final_response = ""
    try:
        for event in runner.run(user_id="user1", session_id="session1", new_message=content):
            if event.is_final_response():
                # Adiciona verificações para content e content.parts
                if event.content is not None and event.content.parts is not None:
                    for part in event.content.parts:
                        if part.text is not None:
                            final_response += part.text
                            final_response += "\n"
                else:
                    # Trata o caso onde content ou content.parts é None
                    error_msg = f"Agente {agent.name} retornou uma resposta final sem conteúdo esperado."
                    print(error_msg) # Para debug no console
                    return f"Erro: {error_msg}"

    except Exception as e:
        # Captura erros durante a execução do agente
        error_msg = f"Erro ao executar o agente {agent.name}: {e}"
        print(error_msg) # Para debug no console
        return f"Erro: {error_msg}"

    if not final_response.strip():
        # Retorna uma mensagem se a resposta final estiver vazia após processar
        return f"Agente {agent.name} não gerou conteúdo para o tópico/entrada fornecida."

    return final_response

# --- Definições dos Agentes ---
# Mantendo as definições dos seus agentes exatamente como estavam

def requisitos(topico):
    """Agente para extrair requisitos do tópico."""
    requisitos = Agent(
        name="requisitos",
        model="gemini-2.5-flash-preview-04-17",
        tools=[google_search],
        instruction="""
        Você é um Agente de Requisitos para um sistema de criação de planos de aula.
        Sua tarefa é ler a descrição fornecida pelo usuário e extrair as seguintes informações específicas de forma clara e estruturada.
        Liste cada item encontrado. Se uma informação não for explicitamente mencionada, indique "Não especificado".
        Itens a extrair:
        - Tema Central
        - Nível de Ensino/Ano (Série)
        - Duração Estimada da Aula
        - Objetivos Preliminares ou Foco Principal (se o usuário mencionar)
        - Contexto da Turma ou Observações Adicionais (se o usuário mencionar algo sobre a turma ou necessidades especiais)
        Apresente o resultado como uma lista ou pares chave-valor.
        """
    )
    entrada_do_requisitos = f"Tópico: {topico}"
    return call_agent(requisitos, entrada_do_requisitos)

def referenciais_curriculares(topico, saida_requisitos):
    """Agente para buscar referenciais curriculares."""
    referenciais_curriculares = Agent(
        name="referenciais_curriculares",
        model="gemini-2.5-flash-preview-04-17",
        instruction="""
        Você é um Integrador de Referenciais Curriculares para planos de aula.
        Você recebeu os requisitos básicos do plano (Tema, Nível de Ensino, etc.).
        Sua tarefa é identificar as diretrizes, habilidades e competências relevantes da BNCC (Base Nacional Comum Curricular), LDB (Lei de Diretrizes e Bases da Educação)
        e, se possível, de currículos estaduais/municipais comuns para este Tema e Nível de Ensino.
        Também é responsável por levantar referências bibliográficas para embasar as aulas. Deve criar uma lista com pelo menos cinco referências bibliográficas.
        Use a ferramenta de busca (Google Search) para encontrar essas informações.
        Concentre-se nas habilidades e competências mais diretamente aplicáveis.
        Apresente os referenciais encontrados de forma clara, listando códigos (se houver) e
        descrições das habilidades/competências relevantes.
        """,
        tools=[google_search]
    )
    entrada_do_referenciais_curriculares = f"Tópico:{topico}\nLançamentos buscados: {saida_requisitos}"
    return call_agent(referenciais_curriculares, entrada_do_referenciais_curriculares)

def elaborador(topico, saida_referenciais):
    """Agente para elaborar a estrutura base do plano."""
    elaborador = Agent(
        name="elaborador",
        model="gemini-2.5-flash-preview-04-17",
        instruction="""
        Você é um Elaborador de Plano de Aula Base.
        Você recebeu os requisitos do plano (Tema, Nível, Duração, etc.) e os referenciais curriculares aplicáveis (BNCC, LDB, etc.).
        Sua tarefa é rascunhar a estrutura inicial do plano de aula.
        Inclua as seguintes seções:
        - Título da Aula: Um título claro baseado no tema.
        - Objetivos de Aprendizagem: Escreva objetivos claros e mensuráveis para os alunos, alinhados com os referenciais curriculares fornecidos e o tema/nível.
        - Conteúdo Programático: Liste os principais tópicos ou conceitos que serão abordados na aula.
        - Etapas da Aula: Divida a aula em fases básicas (ex: Introdução/Aquecimento, Desenvolvimento, Conclusão/Fechamento), indicando brevemente o propósito de cada fase.
        Apresente esta estrutura de plano base de forma organizada.
        Também é preciso utilizar o (google_search) para levantar outras cinco referências bibliográficas, para embasar tudo o que foi escrito.
        """,
        tools=[google_search]
    )
    entrada_do_elaborador = f"Tópico:{topico}\nLançamentos buscados: {saida_referenciais}"
    return call_agent(elaborador, entrada_do_elaborador)

def metodologia(topico, saida_elaborador):
    """Agente para detalhar metodologia, atividades e recursos."""
    metodologia = Agent(
        name="metodologia",
        model="gemini-2.5-flash-preview-04-17",
        instruction="""
        Você é um Detalhador de Planos de Aula. Você recebeu um rascunho de plano de aula base (com Título, Objetivos, Conteúdo e Etapas). Sua tarefa é adicionar detalhes práticos e pedagógicos a este rascunho.
        Para cada etapa da aula (Introdução, Desenvolvimento, Conclusão), descreva Sugestões de Atividades específicas que os alunos podem realizar.
        Sugira Metodologias de Ensino apropriadas para o tema, nível e atividades propostas (ex: aula expositiva dialogada, trabalho em grupo, debate, experimentação, uso de mídias).
        Liste os Recursos Didáticos e Materiais necessários para as atividades e metodologias sugeridas (ex: quadro branco, projetor, materiais de experimento, textos, vídeos, online tools).
        Certifique-se de que as sugestões de atividades, metodologias e recursos estejam alinhadas com os Objetivos e o Conteúdo do plano base. Use a ferramenta de busca se precisar de ideias.
        Apresente o plano base expandido com as novas seções e detalhes.
        """,
        tools=[google_search]
    )
    entrada_do_metodologia = f"Tópico:{topico}\nLançamentos buscados: {saida_elaborador}"
    return call_agent(metodologia, entrada_do_metodologia)

def avaliador(topico, saida_metodologia):
    """Agente para sugerir formas de avaliação."""
    # Nota: A instrução deste agente parece ser a mesma do agente 'metodologia'.
    # Mantendo como estava, mas pode ser um ponto a revisar na lógica dos agentes.
    avaliador = Agent(
        name="avaliador",
        model="gemini-2.5-flash-preview-04-17",
        instruction="""
        Você é um Detalhador de Planos de Aula. Você recebeu um rascunho de plano de aula base (com Título, Objetivos, Conteúdo e Etapas). Sua tarefa é adicionar detalhes práticos e pedagógicos a este rascunho.
        Para cada etapa da aula (Introdução, Desenvolvimento, Conclusão), descreva Sugestões de Atividades específicas que os alunos podem realizar.
        Sugira Metodologias de Ensino apropriadas para o tema, nível e atividades propostas (ex: aula expositiva dialogada, trabalho em grupo, debate, experimentação, uso de mídias).
        Liste os Recursos Didáticos e Materiais necessários para as atividades e metodologias sugeridas (ex: quadro branco, projetor, materiais de experimento, textos, vídeos, online tools).
        Certifique-se de que as sugestões de atividades, metodologias e recursos estejam alinhadas com os Objetivos e o Conteúdo do plano base. Use a ferramenta de busca se precisar de ideias.
        Apresente o plano base expandido com as novas seções e detalhes.
        """,
        tools=[google_search]
    )
    entrada_do_avaliador = f"Tópico:{topico}\nLançamentos buscados: {saida_metodologia}"
    return call_agent(avaliador, entrada_do_avaliador)

def coordenador(topico, saida_avaliador):
    """Agente para revisar e sugerir melhorias no plano."""
    coordenador = Agent(
        name="coordenador",
        model="gemini-2.5-flash-preview-04-17",
        instruction="""
        Você é um Coordenador Pedagógico experiente. Você recebeu um rascunho completo de plano de aula. Sua tarefa é revisar este plano de forma crítica e construtiva.
        Verifique a coerência e o alinhamento entre todas as partes: Os objetivos se encaixam nos referenciais? O conteúdo cobre os objetivos? As atividades realmente ajudam a aprender o conteúdo e alcançar os objetivos? A avaliação mede o que foi proposto?
        Avalie a viabilidade prática: O tempo (duração estimada) é realista para todas as atividades e conteúdo? Os recursos são comuns/acessíveis?
        Considere a adequação para o nível de ensino e o contexto da turma (se mencionado nos requisitos).
        Identifique pontos fortes e áreas que precisam de melhoria. Forneça observações pedagógicas e sugestões específicas para refinar o plano. Apresente suas observações e sugestões de forma clara, referenciando as partes do plano a que se aplicam. Você pode adicionar uma nova seção "Observações do Coordenador" ou integrar seus comentários diretamente no texto do plano, marcando-os claramente.
        Confira um por um dos códigos BNCC informados, se estão corretos e sugira a correção dos que estiverem errados.
        """,
        tools=[google_search]
    )
    entrada_do_coordenador = f"Tópico:{topico}\nLançamentos buscados: {saida_avaliador}"
    return call_agent(coordenador, entrada_do_coordenador)

def final(topico, saida_coordenador):
    """Agente para formatar o plano final no template especificado com formatação de texto plano."""
    final = Agent(
        name="final",
        model="gemini-2.5-flash-preview-04-17",
        instruction="""
        Você é o Agente Formatador Final de Planos de Aula. Você receberá um rascunho completo e revisado de um plano de aula. Este rascunho contém diversas informações geradas pelos agentes anteriores, e PODE incluir observações ou comentários de um coordenador pedagógico.

        Sua tarefa é:
        1. Ler e identificar todas as informações relevantes do plano: Título, Tema, Nível/Série, Duração, Requisitos Iniciais, Referenciais Curriculares (incluindo códigos e descrições da BNCC/LDB/etc.), Objetivos de Aprendizagem, Conteúdo Programático, Descrições das Etapas/Atividades, Metodologias, Materiais/Recursos, Sugestões de Avaliação.
        2. **IGNORAR COMPLETAMENTE** quaisquer observações, comentários, sugestões de melhoria ou seções adicionadas pelo "Coordenador Pedagógico" ou outro agente de revisão que não façam parte do conteúdo intrínseco do plano.
        3. **Organizar e formatar SOMENTE o conteúdo relevante do plano** no template exato e com os nomes de seção especificados abaixo.
        4. Garantir que os **códigos da BNCC (e de outros referenciais, se houver)** e suas descrições correspondentes, que foram identificados em etapas anteriores, sejam incluídos na seção "PARÂMETROS CURRICULARES".
        5. **FORMATAR A SAÍDA USANDO APENAS TEXTO PLANO**, adequado para exibição em uma caixa de texto simples. Use MAIÚSCULAS para os títulos das seções e hifens (-) para os itens de listas, com espaçamento adequado para legibilidade.

        O plano final DEVE seguir ESTE TEMPLATE exato e conter APENAS as seções listadas abaixo, nesta ordem, com estes títulos. Use formatação de texto plano como MAIÚSCULAS para títulos e hifens para listas:

        TÍTULO
        [Título claro da aula]

        TURMA/SERIAÇÃO
        [Nível de Ensino/Ano]

        VISÃO GERAL E OBJETIVO
        [Uma breve introdução sobre o tema da aula para esta turma, seguida por uma síntese do objetivo principal ou propósito geral da aula.]

        PARÂMETROS CURRICULARES
        [Liste aqui os códigos completos (ex: EF07CI09) e as descrições das habilidades e competências relevantes da BNCC, LDB ou outros referenciais, conforme identificados na pesquisa inicial. Use hifens para cada item.]

        OBJETIVOS
        [Liste os objetivos de aprendizagem específicos e mensuráveis para os alunos. Use hifens para cada objetivo.]

        CONTEÚDO
        [Liste os tópicos principais ou conceitos que serão abordados na aula. Use hifens para cada tópico.]

        DESENVOLVIMENTO
        [Descreva as etapas da aula em detalhes, incluindo as atividades a serem realizadas em cada fase (Introdução, Desenvolvimento, Conclusão). Transforme as descrições das atividades em um fluxo contínuo ou lista clara de passos da aula. Use hifens para cada passo ou atividade.]

        METODOLOGIA
        [Descreva as principais abordagens e técnicas de ensino que serão utilizadas. Use hifens para cada metodologia.]

        MATERIAIS NECESSÁRIOS
        [Liste todos os recursos didáticos e materiais que serão precisos para a aula. Use hifens para cada item.]

        AVALIAÇÃO
        [Descreva as formas sugeridas para avaliar o aprendizado dos alunos. Use hifens para cada sugestão.]

        REFERÊNCIAS BIBLIOGRÁFICAS
        [Mencione todas as referências bibliográficas utilizadas para elaborar este Planod e Aula. Use hifens para cada referência.]

        Certifique-se de preencher cada seção do template com as informações correspondentes extraídas do input, ignorando o conteúdo do revisor, e usando a formatação de texto plano especificada.
        """,
        tools=[google_search]
    )
    entrada_do_final = f"Tópico:{topico}\nLançamentos buscados: {saida_coordenador}"
    return call_agent(final, entrada_do_final)

# --- Funções para a Interface Gráfica ---

def update_status(message):
    """Atualiza o texto na barra de status da GUI."""
    status_label.config(text=message)
    root.update_idletasks() # Força a atualização da GUI

def generate_plan_threaded():
    """
    Inicia o processo de geração do plano de aula em um thread separado
    para não bloquear a interface gráfica.
    """
    topico = entry_topic.get().strip() # Pega o texto do campo de entrada

    if not topico:
        messagebox.showwarning("Entrada Inválida", "Por favor, digite o tema para criar o plano de aula.")
        return

    # Desabilita a entrada e o botão enquanto processa
    entry_topic.config(state=tk.DISABLED)
    button_generate.config(state=tk.DISABLED)
    # Habilita o botão de exportar apenas após a geração bem-sucedida
    button_export.config(state=tk.DISABLED)

    text_output.config(state=tk.NORMAL) # Habilita para limpar
    text_output.delete(1.0, tk.END) # Limpa a área de texto
    text_output.config(state=tk.DISABLED) # Desabilita novamente

    # Inicia o thread para rodar a lógica dos agentes
    thread = threading.Thread(target=run_agent_sequence, args=(topico,))
    thread.start()

def run_agent_sequence(topico):
    """
    Executa a sequência de chamadas aos agentes.
    Esta função roda no thread separado.
    """
    try:
        update_status("Processando: Requisitos...")
        resultado_req = requisitos(topico)
        # Opcional: exibir resultados intermediários no output
        # append_output("\n--- Resultados (Agente 1: Requisitos) ---\n" + resultado_req)
        if resultado_req.startswith("Erro:"):
            append_output(resultado_req)
            update_status("Erro!")
            return

        update_status("Processando: Referenciais Curriculares...")
        resultado_ref = referenciais_curriculares(topico, resultado_req)
        # append_output("\n--- Resultados (Agente 2: Referenciais Curriculares) ---\n" + resultado_ref)
        if resultado_ref.startswith("Erro:"):
            append_output(resultado_ref)
            update_status("Erro!")
            return

        update_status("Processando: Elaborador...")
        resultado_elab = elaborador(topico, resultado_ref)
        # append_output("\n--- Resultados (Agente 3: Elaborador) ---\n" + resultado_elab)
        if resultado_elab.startswith("Erro:"):
            append_output(resultado_elab)
            update_status("Erro!")
            return

        update_status("Processando: Metodologia...")
        resultado_met = metodologia(topico, resultado_elab)
        # append_output("\n--- Resultados (Agente 4: Metodologia) ---\n" + resultado_met)
        if resultado_met.startswith("Erro:"):
            append_output(resultado_met)
            update_status("Erro!")
            return

        update_status("Processando: Avaliação...")
        resultado_aval = avaliador(topico, resultado_met)
        # append_output("\n--- Resultados (Agente 5: Avaliação) ---\n" + resultado_aval)
        if resultado_aval.startswith("Erro:"):
            append_output(resultado_aval)
            update_status("Erro!")
            return

        update_status("Processando: Revisão (Coordenador)...")
        resultado_coord = coordenador(topico, resultado_aval)
        # append_output("\n--- Resultados (Agente 6: Coordenador) ---\n" + resultado_coord)
        if resultado_coord.startswith("Erro:"):
            append_output(resultado_coord)
            update_status("Erro!")
            return

        update_status("Finalizando: Formatando Plano...")
        resultado_final = final(topico, resultado_coord)
        if resultado_final.startswith("Erro:"):
            append_output(resultado_final)
            update_status("Erro!")
            return

        # Exibe o resultado final na área de texto
        append_output("\n--- Plano de Aula Gerado ---\n" + resultado_final)

        update_status("Processo Concluído!")

        # Habilita o botão de exportar
        button_export.config(state=tk.NORMAL)

    except Exception as e:
        # Captura qualquer erro que ocorra na sequência de agentes
        error_message = f"Ocorreu um erro inesperado durante a geração do plano: {e}"
        append_output("\n" + error_message)
        update_status("Erro!")
        print(f"Erro no thread de processamento: {e}") # Opcional: imprimir no console para debug

    finally:
        # Reabilita a entrada e o botão de gerar ao final do processo (com ou sem erro)
        entry_topic.config(state=tk.NORMAL)
        button_generate.config(state=tk.NORMAL)


def export_to_doc():
    """
    Exporta o conteúdo da área de texto para um arquivo .docx.
    """
    # Pega o conteúdo da área de texto
    plan_text = text_output.get(1.0, tk.END).strip()

    if not plan_text or "--- Plano de Aula Gerado ---" not in plan_text:
        messagebox.showwarning("Nada para Exportar", "Nenhum plano de aula foi gerado ainda.")
        return

    # Abre a caixa de diálogo para salvar o arquivo
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")],
        title="Salvar Plano de Aula como Documento Word"
    )

    if not file_path:
        return # Usuário cancelou

    try:
        document = Document()

        # Adiciona um título principal
        document.add_heading('Plano de Aula', 0)
        document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Processa o texto linha por linha para formatar no Word
        lines = plan_text.split('\n')
        is_list = False
        for line in lines:
            line = line.strip()
            if not line:
                # Adiciona um parágrafo vazio para espaçamento
                document.add_paragraph()
                is_list = False # Reseta o estado de lista
                continue

            if line.startswith("---"): # Ignora as linhas de separação
                is_list = False
                continue

            # Tenta identificar títulos de seção (linhas em maiúsculas)
            # Adicionei uma verificação para evitar que linhas curtas em maiúsculas (como siglas)
            # sejam tratadas como títulos principais, a menos que sejam as seções esperadas.
            # Ajuste esta lógica se os títulos gerados variarem muito.
            if line.isupper() and len(line) > 1 and line in ["TÍTULO", "TURMA/SERIAÇÃO", "VISÃO GERAL E OBJETIVO", "PARÂMETROS CURRICULARES", "OBJETIVOS", "CONTEÚDO", "DESENVOLVIMENTO", "METODOLOGIA", "MATERIAIS NECESSÁRIOS", "AVALIAÇÃO", "REFERÊNCIAS BIBLIOGRÁFICAS"]:
                 document.add_heading(line, level=1) # Adiciona como Título 1 no Word
                 is_list = False
            elif line.startswith('- '): # Identifica itens de lista
                # Se não estamos em uma lista, cria uma nova lista
                if not is_list:
                     paragraph = document.add_paragraph(style='List Bullet')
                     paragraph.add_run(line[2:]) # Adiciona o texto após o '- '
                     is_list = True
                else:
                    # Se já estamos em uma lista, adiciona o item ao último parágrafo da lista
                    paragraph = document.add_paragraph(style='List Bullet')
                    paragraph.add_run(line[2:])
            else: # Trata como parágrafo normal
                document.add_paragraph(line)
                is_list = False # Reseta o estado de lista

        document.save(file_path)
        messagebox.showinfo("Exportação Concluída", f"Plano de aula exportado com sucesso para:\n{file_path}")

    except Exception as e:
        messagebox.showerror("Erro na Exportação", f"Ocorreu um erro ao exportar o documento:\n{e}")


def append_output(text):
    """Adiciona texto à área de saída da GUI."""
    text_output.config(state=tk.NORMAL) # Habilita para adicionar texto
    text_output.insert(tk.END, text)
    text_output.see(tk.END) # Rola para o final
    text_output.config(state=tk.DISABLED) # Desabilita novamente

# --- Configuração da Janela Principal da GUI ---
root = tk.Tk()
root.title("Gerador de Plano de Aula - Kleber Klaar")
root.geometry("1200x720") # Define um tamanho inicial para a janela

# Frame para a entrada do tópico e botões
frame_controls = tk.Frame(root, padx=10, pady=10)
frame_controls.pack(pady=5, fill=tk.X)

label_topic = tk.Label(frame_controls, text="Digite o tema, a turma e demais informações para gerar o plano de aula:", font=('Arial', 12))
label_topic.pack(side=tk.LEFT, padx=(0, 10))

entry_topic = tk.Entry(frame_controls, width=40, font=('Arial', 12)) # Reduzindo um pouco a largura
entry_topic.pack(side=tk.LEFT, expand=True, fill=tk.X)

button_generate = tk.Button(frame_controls, text="Gerar Plano", command=generate_plan_threaded, font=('Arial', 12))
button_generate.pack(side=tk.LEFT, padx=(10, 5))

button_export = tk.Button(frame_controls, text="Exportar para .doc", command=export_to_doc, font=('Arial', 12), state=tk.DISABLED) # Começa desabilitado
button_export.pack(side=tk.LEFT)


# Área de texto para exibir o resultado
# Usamos ScrolledText para ter uma barra de rolagem automática
text_output = scrolledtext.ScrolledText(root, wrap=tk.WORD, state=tk.DISABLED, font=('Arial', 10))
text_output.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

# Barra de status
status_label = tk.Label(root, text="Pronto. Digite o tema acima.", bd=1, relief=tk.SUNKEN, anchor=tk.W, font=('Arial', 10))
status_label.pack(side=tk.BOTTOM, fill=tk.X)

# Inicia o loop principal da interface gráfica
root.mainloop()
